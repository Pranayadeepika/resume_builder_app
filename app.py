from flask import Flask, render_template, request, send_file, url_for
import os
from werkzeug.utils import secure_filename
from xhtml2pdf import pisa
from io import BytesIO
from docx import Document
from docx.shared import Inches

app = Flask(__name__)

UPLOAD_FOLDER = 'static/uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

resume_data = {}

@app.route('/')
def index():
    return render_template('form.html')

@app.route('/resume', methods=['POST'])
def resume():
    global resume_data
    photo_file = request.files.get('photo')
    sign_file = request.files.get('sign_image')

    photo_url = ''
    sign_url = ''
    
    if photo_file and photo_file.filename:
        photo_name = secure_filename(photo_file.filename)
        photo_path = os.path.join(app.config['UPLOAD_FOLDER'], photo_name)
        photo_file.save(photo_path)
        photo_url = url_for('static', filename=f'uploads/{photo_name}')

    if sign_file and sign_file.filename:
        sign_name = secure_filename(sign_file.filename)
        sign_path = os.path.join(app.config['UPLOAD_FOLDER'], sign_name)
        sign_file.save(sign_path)
        sign_url = url_for('static', filename=f'uploads/{sign_name}')

    resume_data = {
        'name': request.form['name'],
        'email': request.form['email'],
        'phone': request.form['phone'],
        'address': request.form['address'],
        'linkedin': request.form['linkedin'],
        'photo': photo_url,
        'sign_image': sign_url,
        'education': request.form['education'].split('\n'),
        'projects': request.form['projects'].split('\n'),
        'courses': request.form['courses'].split('\n'),
        'skills': request.form['skills'].split(','),
        'languages': request.form['languages'].split(','),
        'hobbies': request.form['hobbies'].split(','),
        'dob': request.form['dob'],
        'gender': request.form['gender']
    }
    return render_template('resume_template.html', **resume_data)

@app.route('/download/pdf')
def download_pdf():
    rendered = render_template('resume_template.html', **resume_data)
    result = BytesIO()
    pisa.CreatePDF(rendered, dest=result)
    result.seek(0)
    return send_file(result, as_attachment=True, download_name='resume.pdf', mimetype='application/pdf')

@app.route('/download/docx')
def download_docx():
    doc = Document()
    doc.add_heading(resume_data['name'], 0)
    doc.add_paragraph(f"Email: {resume_data['email']} | Phone: {resume_data['phone']}")
    doc.add_paragraph(f"Address: {resume_data['address']}")
    doc.add_paragraph(f"LinkedIn: {resume_data['linkedin']}")

    doc.add_heading("Education", level=1)
    for edu in resume_data['education']:
        doc.add_paragraph(edu, style='List Bullet')

    doc.add_heading("Projects", level=1)
    for p in resume_data['projects']:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading("Courses / Certifications", level=1)
    for c in resume_data['courses']:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_heading("Skills", level=1)
    for s in resume_data['skills']:
        doc.add_paragraph(s.strip(), style='List Bullet')

    doc.add_heading("Languages Known", level=1)
    for l in resume_data['languages']:
        doc.add_paragraph(l.strip(), style='List Bullet')

    doc.add_heading("Hobbies / Interests", level=1)
    for h in resume_data['hobbies']:
        doc.add_paragraph(h.strip(), style='List Bullet')

    doc.add_heading("Personal Details", level=1)
    doc.add_paragraph(f"Date of Birth: {resume_data['dob']}")
    doc.add_paragraph(f"Gender: {resume_data['gender']}")

    result = BytesIO()
    doc.save(result)
    result.seek(0)
    return send_file(result, as_attachment=True, download_name='resume.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# ✅ Port binding for Render deployment
if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
